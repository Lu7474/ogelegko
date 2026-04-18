import io
import logging
import threading
import uuid
import zipfile

from django.core.cache import cache
from django.db import IntegrityError
from django.db.models import Count, Q
from django.http import HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.views.decorators.http import require_POST

from .admin_views import _safe_int, admin_required
from .models import Answer, Attempt, ExamType, Task, TaskSource, Variant
from .parser import sanitize_html

logger = logging.getLogger(__name__)

_JOB_TTL = 7200  # 2 часа


# ===== ВАРИАНТЫ =====


@admin_required
def variant_list(request):
    exam_filter = request.GET.get("exam_type", "oge")
    variants = Variant.objects.annotate(
        task_count=Count("tasks", distinct=True),
        attempts_count=Count("attempts", filter=Q(attempts__is_finished=True), distinct=True),
    )
    if exam_filter and exam_filter in dict(ExamType.choices):
        variants = variants.filter(exam_type=exam_filter)

    return render(
        request,
        "admin/variants.html",
        {
            "variants": variants,
            "exam_types": ExamType.choices,
            "exam_filter": exam_filter,
        },
    )


def _save_variant_tasks(variant, request):
    """Сохраняет задания варианта из POST-данных (для новых вариантов)."""
    import re as _re

    indices = sorted(
        {int(m.group(1)) for key in request.POST for m in [_re.match(r"^task_(\d+)_answer$", key)] if m}
    )
    for task_index in indices:
        text = sanitize_html(request.POST.get(f"task_{task_index}_text", "").strip())
        answer = request.POST.get(f"task_{task_index}_answer", "").strip()
        source = request.POST.get(f"task_{task_index}_source", "manual")
        points = _safe_int(request.POST.get(f"task_{task_index}_points", "1"), default=1)
        manual_grading = bool(request.POST.get(f"task_{task_index}_manual_grading"))
        image = request.FILES.get(f"task_{task_index}_image")
        shared_context = sanitize_html(request.POST.get(f"task_{task_index}_shared_context", "").strip())
        shared_context_image = request.FILES.get(f"task_{task_index}_shared_context_image")

        if points < 1:
            points = 1
        if source not in dict(TaskSource.choices):
            source = "manual"

        if answer or manual_grading:
            task = Task(
                variant=variant,
                number=task_index,
                text=text,
                correct_answer=answer,
                source=source,
                points=points,
                manual_grading=manual_grading,
                shared_context=shared_context,
            )
            if image:
                task.image = image
            if shared_context_image:
                task.shared_context_image = shared_context_image
            task.save()


def _update_variant_tasks(variant, request):
    """Обновляет задания существующего варианта без удаления (сохраняет ответы учеников)."""
    import re as _re

    indices = sorted(
        {int(m.group(1)) for key in request.POST for m in [_re.match(r"^task_(\d+)_answer$", key)] if m}
    )
    task_numbers_in_form = set()
    for task_index in indices:
        text = sanitize_html(request.POST.get(f"task_{task_index}_text", "").strip())
        answer = request.POST.get(f"task_{task_index}_answer", "").strip()
        source = request.POST.get(f"task_{task_index}_source", "manual")
        points = _safe_int(request.POST.get(f"task_{task_index}_points", "1"), default=1)
        manual_grading = bool(request.POST.get(f"task_{task_index}_manual_grading"))
        image = request.FILES.get(f"task_{task_index}_image")
        shared_context = sanitize_html(request.POST.get(f"task_{task_index}_shared_context", "").strip())
        shared_context_image = request.FILES.get(f"task_{task_index}_shared_context_image")

        if points < 1:
            points = 1
        if source not in dict(TaskSource.choices):
            source = "manual"

        if answer or manual_grading:
            number_str = str(task_index)
            task_numbers_in_form.add(number_str)
            task, _ = Task.objects.update_or_create(
                variant=variant,
                number=number_str,
                defaults={
                    "text": text,
                    "correct_answer": answer,
                    "source": source,
                    "points": points,
                    "manual_grading": manual_grading,
                    "shared_context": shared_context,
                },
            )
            if image:
                task.image = image
                task.save(update_fields=["image"])
            if shared_context_image:
                task.shared_context_image = shared_context_image
                task.save(update_fields=["shared_context_image"])

    # Удалить задания, которых больше нет в форме (каскадно удаляет их ответы)
    variant.tasks.exclude(number__in=task_numbers_in_form).delete()


@admin_required
def variant_add(request):
    error = ""
    if request.method == "POST":
        number = request.POST.get("number", "").strip()
        exam_type = request.POST.get("exam_type", "")
        max_attempts = _safe_int(request.POST.get("max_attempts", "3"), default=3)

        if not number:
            error = "Введите номер варианта"
        elif exam_type not in dict(ExamType.choices):
            error = "Выберите тип экзамена"
        else:
            try:
                variant = Variant.objects.create(
                    number=number, exam_type=exam_type, max_attempts=max_attempts
                )
                _save_variant_tasks(variant, request)
                return redirect("admin_variants")
            except IntegrityError:
                error = f"Вариант с номером '{number}' уже существует"

    return render(
        request,
        "admin/variant_form.html",
        {
            "exam_types": ExamType.choices,
            "sources": TaskSource.choices,
            "error": error,
        },
    )


@admin_required
def variant_edit(request, variant_id):
    variant = get_object_or_404(Variant, id=variant_id)
    tasks = variant.tasks.order_by("id")
    error = ""

    if request.method == "POST":
        number = request.POST.get("number", "").strip()
        exam_type = request.POST.get("exam_type", variant.exam_type)
        max_attempts = _safe_int(request.POST.get("max_attempts", "3"), default=3)

        if not number:
            error = "Введите номер варианта"
        else:
            try:
                variant.number = number
                variant.exam_type = exam_type
                variant.max_attempts = max_attempts
                variant.save()

                _update_variant_tasks(variant, request)
                return redirect("admin_variants")
            except IntegrityError:
                error = f"Вариант с номером '{number}' уже существует"

    return render(
        request,
        "admin/variant_form.html",
        {
            "variant": variant,
            "tasks": tasks,
            "exam_types": ExamType.choices,
            "sources": TaskSource.choices,
            "error": error,
        },
    )


@admin_required
@require_POST
def variant_toggle(request, variant_id):
    variant = get_object_or_404(Variant, id=variant_id)
    variant.is_active = not variant.is_active
    variant.save(update_fields=["is_active"])
    return redirect("admin_variants")


@admin_required
@require_POST
def variant_duplicate(request, variant_id):
    original = get_object_or_404(Variant, id=variant_id)
    copy_number = f"{original.number}_копия"
    counter = 1
    while Variant.objects.filter(number=copy_number).exists():
        counter += 1
        copy_number = f"{original.number}_копия{counter}"

    new_variant = Variant.objects.create(
        number=copy_number,
        exam_type=original.exam_type,
        max_attempts=original.max_attempts,
    )
    for task in original.tasks.all():
        Task.objects.create(
            variant=new_variant,
            number=task.number,
            text=task.text,
            image=task.image,
            correct_answer=task.correct_answer,
            source=task.source,
            points=task.points,
            manual_grading=task.manual_grading,
            no_student_input=task.no_student_input,
            shared_context=task.shared_context,
            shared_context_image=task.shared_context_image,
        )
    return redirect("admin_variant_edit", variant_id=new_variant.id)


@admin_required
@require_POST
def variant_delete(request, variant_id):
    variant = get_object_or_404(Variant, id=variant_id)
    variant.delete()
    return redirect("admin_variants")


# ===== ИМПОРТ ВАРИАНТА =====


def _run_import_job(job_id, url, variant_number):
    """Фоновый поток для импорта варианта."""
    try:
        from .parser import import_variant_from_sdamgia

        variant, parse_errors = import_variant_from_sdamgia(url, variant_number=variant_number or None)
        cache.set(
            f"vjob:{job_id}",
            {
                "status": "done",
                "variant_id": variant.id if variant else None,
                "errors": parse_errors,
            },
            _JOB_TTL,
        )
    except Exception as e:
        logger.exception("Ошибка импорта варианта")
        cache.set(
            f"vjob:{job_id}",
            {"status": "error", "variant_id": None, "errors": [str(e)]},
            _JOB_TTL,
        )
    finally:
        from django.db import connection

        connection.close()


@admin_required
def variant_import(request):
    """Импорт варианта с sdamgia.ru."""
    if request.method == "POST":
        url = request.POST.get("url", "").strip()
        variant_number = request.POST.get("variant_number", "").strip()

        errors = []
        if not url:
            errors.append("Введите URL варианта")
        elif "sdamgia.ru" not in url:
            errors.append("Поддерживается только sdamgia.ru (Решу ОГЭ / Решу ЕГЭ)")

        if errors:
            return render(request, "admin/variant_import.html", {"errors": errors})

        job_id = str(uuid.uuid4())
        cache.set(f"vjob:{job_id}", {"status": "running", "variant_id": None, "errors": []}, _JOB_TTL)
        threading.Thread(target=_run_import_job, args=(job_id, url, variant_number), daemon=True).start()
        return redirect("admin_variant_import_status", job_id=job_id)

    return render(request, "admin/variant_import.html", {"errors": []})


@admin_required
def variant_import_status(request, job_id):
    """Страница/API опроса статуса импорта."""
    job = cache.get(f"vjob:{job_id}") or {
        "status": "unknown",
        "variant_id": None,
        "errors": ["Задание не найдено"],
    }

    if request.GET.get("json"):
        return JsonResponse(job)

    variant = None
    if job["status"] == "done" and job.get("variant_id"):
        variant = Variant.objects.filter(id=job["variant_id"]).first()

    return render(
        request,
        "admin/variant_import_status.html",
        {"job_id": job_id, "job": job, "variant": variant},
    )


@admin_required
def variant_stats(request, variant_id):
    variant = get_object_or_404(Variant, id=variant_id)
    attempts = (
        Attempt.objects.filter(variant=variant, is_finished=True)
        .select_related("student", "student__school_class")
        .order_by("-finished_at")
    )

    avg_percentage = None
    if attempts.exists():
        percentages = [a.percentage for a in attempts]
        avg_percentage = round(sum(percentages) / len(percentages))

    task_stats = []
    for task in variant.tasks.order_by("id"):
        total = Answer.objects.filter(task=task, attempt__is_finished=True).count()
        correct = Answer.objects.filter(task=task, attempt__is_finished=True, is_correct=True).count()
        pct = round(correct / total * 100) if total > 0 else 0
        task_stats.append({"task": task, "correct": correct, "total": total, "percentage": pct})

    return render(
        request,
        "admin/variant_stats.html",
        {
            "variant": variant,
            "attempts": attempts,
            "avg_percentage": avg_percentage,
            "task_stats": task_stats,
        },
    )


# ===== ПЕЧАТЬ ВАРИАНТА (DOCX) =====


def _parse_html_segments(html):
    """Разбирает HTML на сегменты: ('text',...) | ('image', src) | ('break',) | ('table', rows)."""
    from bs4 import BeautifulSoup, NavigableString, Tag

    segments = []

    from bs4 import Comment

    def walk(node, bold=False, italic=False, sup=False, sub=False):
        if isinstance(node, Comment):
            return  # пропускаем HTML-комментарии
        if isinstance(node, NavigableString):
            text = str(node).replace("\r", "").replace("\n", " ")
            if text:
                segments.append(("text", text, bold, italic, sup, sub))
            return
        if not isinstance(node, Tag):
            return
        tag = node.name.lower()
        if tag == "img":
            src = node.get("src", "")
            if src:
                style = node.get("style", "")
                if "float:right" in style or "float: right" in style:
                    segments.append(("image_right", src))
                else:
                    segments.append(("image", src))
            return
        if tag in ("details", "summary"):
            # <details class="shared-context"> извлекается до вызова этой функции
            return
        if tag == "br":
            segments.append(("break",))
            return
        if tag in ("p", "div", "li"):
            if segments and segments[-1][0] != "break":
                segments.append(("break",))
            for c in node.children:
                walk(c, bold, italic, sup, sub)
            if segments and segments[-1][0] != "break":
                segments.append(("break",))
            return
        if tag == "table":
            rows = []
            for tr in node.find_all("tr"):
                cells = [str(td) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            if rows:
                segments.append(("table", rows))
            return
        if tag in ("td", "th"):
            for c in node.children:
                walk(c, bold, italic, sup, sub)
            segments.append(("text", "  ", False, False, False, False))
            return
        if tag == "tr":
            for c in node.children:
                walk(c, bold, italic, sup, sub)
            segments.append(("break",))
            return
        if tag in ("b", "strong"):
            for c in node.children:
                walk(c, True, italic, sup, sub)
        elif tag in ("i", "em"):
            for c in node.children:
                walk(c, bold, True, sup, sub)
        elif tag == "sup":
            for c in node.children:
                walk(c, bold, italic, True, False)
        elif tag == "sub":
            for c in node.children:
                walk(c, bold, italic, False, True)
        elif tag == "span":
            classes = node.get("class", [])
            if isinstance(classes, str):
                classes = classes.split()
            if "math-frac" in classes:
                spans = list(node.find_all("span", recursive=False))
                if len(spans) >= 2:
                    num = spans[0].get_text(strip=True)
                    den = spans[1].get_text(strip=True)
                    segments.append(("text", f"({num})/({den})", bold, italic, sup, sub))
                else:
                    for c in node.children:
                        walk(c, bold, italic, sup, sub)
            else:
                for c in node.children:
                    walk(c, bold, italic, sup, sub)
        else:
            for c in node.children:
                walk(c, bold, italic, sup, sub)

    soup = BeautifulSoup(html or "", "html.parser")
    for child in soup.children:
        walk(child)
    while segments and segments[-1][0] == "break":
        segments.pop()

    # Переставляем image_right в конец каждой группы (до break),
    # чтобы картинка шла после текста, а не перед ним.
    reordered, group_text, group_imgs = [], [], []
    for seg in segments:
        if seg[0] == "break":
            reordered.extend(group_text)
            reordered.extend(group_imgs)
            reordered.append(seg)
            group_text, group_imgs = [], []
        elif seg[0] == "image_right":
            group_imgs.append(seg)
        else:
            group_text.append(seg)
    reordered.extend(group_text)
    reordered.extend(group_imgs)
    return reordered


def _get_image_bytes(src):
    """Загружает изображение по src (/media/... или http...) и возвращает bytes или None."""
    import os

    import requests as _req
    from django.conf import settings as dj_settings

    try:
        if src.startswith("http://") or src.startswith("https://"):
            r = _req.get(src, timeout=15)
            if r.status_code == 200:
                return r.content
        elif src.startswith("/media/"):
            rel = src[len("/media/") :]
            local = os.path.join(str(getattr(dj_settings, "MEDIA_ROOT", "")), rel.replace("/", os.sep))
            if os.path.exists(local):
                with open(local, "rb") as f:
                    return f.read()
            from django.core.files.storage import default_storage

            try:
                url = default_storage.url(rel)
                if url.startswith("http://") or url.startswith("https://"):
                    r = _req.get(url, timeout=15)
                    if r.status_code == 200:
                        return r.content
                else:
                    with default_storage.open(rel) as f:
                        return f.read()
            except Exception:
                pass
    except Exception as e:
        logger.warning("Не удалось загрузить изображение %s: %s", src, e)
    return None


def _svg_to_png(svg_bytes):
    """Конвертирует SVG байты в PNG байты через PyMuPDF. Возвращает None при ошибке."""
    try:
        import fitz  # PyMuPDF

        fitz_doc = fitz.open("svg", svg_bytes)
        pix = fitz_doc[0].get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
        return pix.tobytes("png")
    except Exception as e:
        logger.warning("SVG→PNG конвертация не удалась: %s", e)
        return None


def _image_width(img_data, max_cm=14):
    """Возвращает ширину для вставки изображения: натуральный размер, но не более max_cm."""
    from docx.shared import Cm

    try:
        from PIL import Image as PILImage

        img = PILImage.open(io.BytesIO(img_data))
        w_px = img.size[0]
        info_dpi = img.info.get("dpi", None)
        if info_dpi:
            dpi_x = info_dpi[0] if hasattr(info_dpi, "__getitem__") else info_dpi
        else:
            dpi_x = 96
        if not dpi_x or dpi_x <= 0:
            dpi_x = 96
        w_cm = w_px / dpi_x * 2.54
        return Cm(min(w_cm, max_cm))
    except Exception:
        return Cm(max_cm)


def _render_segments(doc, segments, indent=None, font_size=None, initial_para=None, default_font=None):
    """Рендерит сегменты в документ, создавая параграфы по мере нужды."""
    from docx.shared import Pt

    current = [initial_para]

    def get_para():
        if current[0] is None:
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.left_indent = indent
            p.paragraph_format.space_after = Pt(1)
            current[0] = p
        return current[0]

    def close():
        current[0] = None

    def _apply(run):
        pass  # font inherited from Normal style

    for seg in segments:
        if seg[0] == "text":
            _, text, bold, italic, sup, sub = seg
            run = get_para().add_run(text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            run.font.superscript = sup
            run.font.subscript = sub
            _apply(run)
        elif seg[0] == "break":
            close()
        elif seg[0] in ("image", "image_right"):
            img_src = seg[1]
            is_right = seg[0] == "image_right"
            img_data = _get_image_bytes(img_src)
            if img_data:
                is_svg = img_src.lower().endswith(".svg") or img_data[:5] in (b"<svg ", b"<?xml")
                if is_svg:
                    img_data = _svg_to_png(img_data)
                if img_data:
                    try:
                        if is_svg:
                            # SVG inline — до 5 см (формулы и чертежи одинаково).
                            max_svg = 5
                            if is_right:
                                max_svg = min(max_svg, 3)
                            get_para().add_run().add_picture(
                                io.BytesIO(img_data), width=_image_width(img_data, max_cm=max_svg)
                            )
                        elif is_right:
                            # float:right PNG/JPEG — компактное (до 3 см)
                            get_para().add_run().add_picture(
                                io.BytesIO(img_data), width=_image_width(img_data, max_cm=3)
                            )
                        else:
                            # Обычное изображение — отдельным блоком (до 4 см)
                            close()
                            doc.add_picture(io.BytesIO(img_data), width=_image_width(img_data, max_cm=4))
                            close()
                    except Exception as e:
                        logger.warning("Не удалось вставить изображение: %s", e)
        elif seg[0] == "table":
            close()
            rows_data = seg[1]
            if not rows_data:
                continue
            max_cols = max(len(row) for row in rows_data)
            if max_cols == 0:
                continue
            tbl = doc.add_table(rows=len(rows_data), cols=max_cols)
            tbl.style = "Table Grid"
            for r_idx, row in enumerate(rows_data):
                for c_idx, cell_html in enumerate(row):
                    if c_idx >= max_cols:
                        break
                    cell = tbl.rows[r_idx].cells[c_idx]
                    cell_segs = _parse_html_segments(cell_html)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    for seg2 in cell_segs:
                        if seg2[0] == "text":
                            run = p.add_run(seg2[1])
                            if seg2[2]:
                                run.bold = True
                            if seg2[3]:
                                run.italic = True
                            run.font.superscript = seg2[4]
                            run.font.subscript = seg2[5]
                            _apply(run)
                        elif seg2[0] == "break":
                            p = cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(1)
            close()


def _extract_task_parts(html):
    """Делит task.text на (ctx_html, body_html).

    Ищет <details class="shared-context"> → возвращает содержимое .shared-context-body
    как ctx_html (или None если нет <details>), а остаток без <details> — как body_html.
    """
    import re

    from bs4 import BeautifulSoup

    if not html or "shared-context" not in html:
        return None, html

    soup = BeautifulSoup(html, "html.parser")
    details = soup.find("details", class_="shared-context")
    if not details:
        return None, html

    body_div = details.find("div", class_="shared-context-body")
    ctx_html = str(body_div) if body_div else ""
    details.decompose()

    body_html = re.sub(r"^(\s*<br\s*/?>)+", "", str(soup)).strip()
    return ctx_html, body_html


def _strip_answer_placeholder(html):
    """Удаляет строки 'Ответ: ___' из HTML задания."""
    import re

    from bs4 import BeautifulSoup

    if not html:
        return html
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all(["p", "div", "span"]):
        text = tag.get_text()
        if re.search(r"Ответ\s*:\s*[_\s]{3,}", text):
            tag.decompose()
    return str(soup)


def _build_variant_docx(variant, include_answers):
    """Строит docx-документ для варианта. Возвращает BytesIO.

    Формат: US Letter, узкие поля — параметры эталонного «33 вариант.docx».
    """
    import requests as _req
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Inches, Pt

    doc = Document()

    # Шрифт: Calibri 11pt — соответствует эталону (нет явного шрифта → тема Word = Calibri,
    # sz=22 в docDefaults = 11pt)
    from docx.oxml.ns import qn as _qn

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    rPr = style_normal.element.get_or_add_rPr()
    rFonts = rPr.find(_qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement as _OE

        rFonts = _OE("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(_qn("w:ascii"), "Calibri")
    rFonts.set(_qn("w:hAnsi"), "Calibri")
    rFonts.set(_qn("w:cs"), "Calibri")

    # Размер страницы US Letter + поля — как в эталоне
    # (12240 × 15840 twips = 8.5 × 11 дюймов)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    FS = Pt(11)
    FONT = "Calibri"

    def _sp(p, before=0, after=1):
        if before:
            p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)

    def _run(para, text, bold=False, size=None):
        """Добавляет run; size ставится только если передан явно."""
        r = para.add_run(text)
        if size:
            r.font.size = size
        if bold:
            r.bold = True
        return r

    def _shade_cell(cell, fill="D9D9D9"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    # ─── Заголовок варианта ──────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(title, before=0, after=1)
    _run(title, f"Вариант {variant.number}", bold=True, size=Pt(13))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(sub, before=0, after=3)
    _run(sub, variant.get_exam_type_display(), size=Pt(10))

    tasks = sorted(
        variant.tasks.all(),
        key=lambda t: int(t.number) if t.number.isdigit() else 0,
    )

    # Хэши уже напечатанных общих условий (дедупликация)
    printed_ctx_hashes: set = set()
    part1_printed = False
    part2_printed = False

    for task in tasks:
        task_num = int(task.number) if task.number.isdigit() else 0

        # ─── Заголовок "Часть 1." / "Часть 2." ──────────────────────────
        if task_num >= 20 and not part2_printed:
            part2_printed = True
            p2 = doc.add_paragraph()
            _sp(p2, before=6, after=2)
            _run(p2, "Часть 2.", bold=True)
        elif not part1_printed and task_num < 20:
            part1_printed = True
            p1 = doc.add_paragraph()
            _sp(p1, before=0, after=2)
            _run(p1, "Часть 1.", bold=True)

        # ─── Общее условие из task.shared_context (legacy) ──────────────
        if task.shared_context or task.shared_context_image:
            ctx_key = hash(task.shared_context or "")
            if ctx_key not in printed_ctx_hashes:
                printed_ctx_hashes.add(ctx_key)
                if task.shared_context_image:
                    try:
                        ci_url = task.shared_context_image.url
                        ci_data = (
                            _req.get(ci_url, timeout=15).content
                            if ci_url.startswith("http")
                            else task.shared_context_image.open("rb").read()
                        )
                        doc.add_picture(io.BytesIO(ci_data), width=_image_width(ci_data, max_cm=12))
                    except Exception:
                        logger.warning("Не удалось вставить изображение общего условия")
                if task.shared_context:
                    _render_segments(
                        doc,
                        _parse_html_segments(task.shared_context),
                        font_size=FS,
                        default_font=FONT,
                    )

        # ─── Общее условие из <details class="shared-context"> в task.text ─
        ctx_html, body_html = _extract_task_parts(task.text or "")
        if ctx_html:
            ctx_key = hash(ctx_html.strip())
            if ctx_key not in printed_ctx_hashes:
                printed_ctx_hashes.add(ctx_key)
                _render_segments(
                    doc,
                    _parse_html_segments(ctx_html),
                    font_size=FS,
                    default_font=FONT,
                )

        # ─── Номер задания + текст в одном параграфе ────────────────────
        th = doc.add_paragraph()
        _sp(th, before=3, after=1)
        _run(th, f"{task.number}. ", bold=True)

        # ─── Текст задания (без «Ответ: ___»), продолжает th ────────────
        body_html = _strip_answer_placeholder(body_html)
        if body_html.strip():
            _render_segments(
                doc,
                _parse_html_segments(body_html),
                font_size=FS,
                default_font=FONT,
                initial_para=th,
            )

        # ─── Картинка задания (legacy) ───────────────────────────────────
        if task.image:
            try:
                img_url = task.image.url
                img_data = (
                    _req.get(img_url, timeout=15).content
                    if img_url.startswith("http")
                    else task.image.open("rb").read()
                )
                doc.add_picture(io.BytesIO(img_data), width=_image_width(img_data))
            except Exception:
                logger.warning("Не удалось вставить картинку задания %s", task.number)

    # ─── Таблица ответов (только для учителя, на отдельной странице) ────
    auto_tasks = [t for t in tasks if int(t.number) < 20 if t.number.isdigit()]
    if not auto_tasks:
        auto_tasks = tasks  # fallback если номера нестандартные
    if auto_tasks and include_answers:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        sep = doc.add_paragraph()
        _sp(sep, before=0, after=6)
        _run(sep, "Таблица ответов", bold=True, size=Pt(13))

        col_groups = 1
        n = len(auto_tasks)
        rpg = (n + col_groups - 1) // col_groups
        groups = [auto_tasks[g * rpg : (g + 1) * rpg] for g in range(col_groups)]

        tbl = doc.add_table(rows=rpg + 1, cols=col_groups * 2)
        tbl.style = "Table Grid"
        # Узкие колонки: № — 1 cm, Ответ — 3 cm
        for row in tbl.rows:
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(3)

        hdr = tbl.rows[0].cells
        for g in range(col_groups):
            hdr[g * 2].text = "№"
            hdr[g * 2 + 1].text = "Ответ"

        for row_i in range(rpg):
            cells = tbl.rows[row_i + 1].cells
            for g, grp in enumerate(groups):
                if row_i < len(grp):
                    cells[g * 2].text = str(grp[row_i].number)
                    cells[g * 2 + 1].text = grp[row_i].correct_answer or ""

        for row_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _sp(para, before=1, after=1)
                    for run in para.runs:
                        run.font.name = FONT
                        run.font.size = Pt(9)
                        if row_idx == 0 or c_idx % 2 == 0:
                            run.bold = True
                if row_idx == 0 or c_idx % 2 == 0:
                    _shade_cell(cell)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@admin_required
def variant_print_docx(request, variant_id, mode):
    """Скачать docx: mode='teacher' (с ответами) или 'student' (без)."""
    from urllib.parse import quote

    variant = get_object_or_404(Variant, id=variant_id)
    safe_num = variant.number.replace("/", "-")
    include_answers = mode == "teacher"
    suffix = "_answers" if include_answers else ""

    buf = _build_variant_docx(variant, include_answers=include_answers)

    fname = f"{safe_num}_variant{suffix}.docx"
    fname_utf8 = quote(f"{safe_num} вариант.docx")
    response = HttpResponse(
        buf,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f"attachment; filename=\"{fname}\"; filename*=UTF-8''{fname_utf8}"
    return response


@admin_required
@require_POST
def variants_print_zip(request):
    """Скачать ZIP с DOCX-файлами для выбранных вариантов."""
    ids = request.POST.getlist("ids")
    if not ids:
        return HttpResponse("Не выбрано ни одного варианта", status=400)

    variants = Variant.objects.filter(id__in=ids)
    if not variants.exists():
        return HttpResponse("Варианты не найдены", status=404)

    from urllib.parse import quote

    numbers = [v.number for v in variants]
    safe_numbers = ", ".join(n.replace("/", "-") for n in numbers)
    zip_name_ascii = f"{safe_numbers}.zip"
    zip_name_utf8 = quote(f"{', '.join(numbers)}.zip")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for variant in variants:
            docx_buf = _build_variant_docx(variant, include_answers=True)
            safe_num = variant.number.replace("/", "-")
            zf.writestr(f"{safe_num} вариант.docx", docx_buf.read())

    buf.seek(0)
    response = HttpResponse(buf, content_type="application/zip")
    response["Content-Disposition"] = (
        f"attachment; filename=\"{zip_name_ascii}\"; filename*=UTF-8''{zip_name_utf8}"
    )
    return response


@admin_required
@require_POST
def variants_archive_export(request):
    """Скачать переносимый ZIP-архив выбранных (или всех) вариантов."""
    from datetime import datetime

    from django.contrib import messages
    from django.http import FileResponse

    from .services.variant_archive import export_variants_to_zip

    ids = request.POST.getlist("ids")
    if ids:
        variants = Variant.objects.filter(id__in=ids)
    else:
        variants = Variant.objects.all()

    if not variants.exists():
        messages.error(request, "Нет вариантов для экспорта.")
        return redirect("admin_variants")

    buf = export_variants_to_zip(variants)
    filename = f"variants_archive_{datetime.now():%Y-%m-%d_%H-%M}.zip"
    return FileResponse(buf, as_attachment=True, filename=filename, content_type="application/zip")


@admin_required
@require_POST
def variants_archive_import(request):
    """Импортировать варианты из переносимого ZIP-архива."""
    from django.contrib import messages

    from .services.variant_archive import ArchiveImportError, import_variants_from_zip

    uploaded = request.FILES.get("archive")
    if not uploaded:
        messages.error(request, "Файл не выбран.")
        return redirect("admin_variants")

    try:
        result = import_variants_from_zip(uploaded)
    except ArchiveImportError as e:
        messages.error(request, f"Ошибка архива: {e}")
        return redirect("admin_variants")
    except Exception as e:
        logger.exception("Непредвиденная ошибка при импорте архива вариантов")
        messages.error(request, f"Ошибка импорта: {e}")
        return redirect("admin_variants")

    parts = [f"Импортировано: {result['variants_created']} вар., {result['tasks_created']} зад."]
    if result["renamed"]:
        parts.append("Переименованы: " + "; ".join(result["renamed"]))
    if result["errors"]:
        parts.append("Ошибки: " + "; ".join(result["errors"]))

    messages.success(request, " ".join(parts))
    return redirect("admin_variants")


@admin_required
@require_POST
def variants_bulk_toggle(request):
    """Массовая активация / скрытие вариантов."""
    ids = request.POST.getlist("ids")
    action = request.POST.get("action")
    if ids and action in ("activate", "hide"):
        Variant.objects.filter(id__in=ids).update(is_active=(action == "activate"))
    return redirect("admin_variants")


@admin_required
@require_POST
def variants_bulk_delete(request):
    """Массовое удаление вариантов."""
    ids = request.POST.getlist("ids")
    if ids:
        Variant.objects.filter(id__in=ids).delete()
    return redirect("admin_variants")
