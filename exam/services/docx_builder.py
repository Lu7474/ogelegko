import io
import logging

logger = logging.getLogger(__name__)


def _parse_html_segments(html):
    """Разбирает HTML на сегменты: ('text',...) | ('image', src) | ('break',) | ('table', rows)."""
    from bs4 import BeautifulSoup, Comment, NavigableString, Tag

    segments = []

    def walk(node, bold=False, italic=False, sup=False, sub=False):
        if isinstance(node, Comment):
            return
        if isinstance(node, NavigableString):
            text = str(node).replace("\r", "").replace("\n", " ")
            if text:
                segments.append(("text", text, bold, italic, sup, sub))
            return
        if not isinstance(node, Tag):
            return
        tag = node.name.lower()
        if tag == "img":
            src = node.get("src", "")
            if src:
                style = node.get("style", "")
                if "float:right" in style or "float: right" in style:
                    segments.append(("image_right", src))
                else:
                    segments.append(("image", src))
            return
        if tag in ("details", "summary"):
            return
        if tag == "br":
            segments.append(("break",))
            return
        if tag in ("p", "div", "li"):
            if segments and segments[-1][0] != "break":
                segments.append(("break",))
            for c in node.children:
                walk(c, bold, italic, sup, sub)
            if segments and segments[-1][0] != "break":
                segments.append(("break",))
            return
        if tag == "table":
            rows = []
            for tr in node.find_all("tr"):
                cells = [str(td) for td in tr.find_all(["td", "th"])]
                if cells:
                    rows.append(cells)
            if rows:
                segments.append(("table", rows))
            return
        if tag in ("td", "th"):
            for c in node.children:
                walk(c, bold, italic, sup, sub)
            segments.append(("text", "  ", False, False, False, False))
            return
        if tag == "tr":
            for c in node.children:
                walk(c, bold, italic, sup, sub)
            segments.append(("break",))
            return
        if tag in ("b", "strong"):
            for c in node.children:
                walk(c, True, italic, sup, sub)
        elif tag in ("i", "em"):
            for c in node.children:
                walk(c, bold, True, sup, sub)
        elif tag == "sup":
            for c in node.children:
                walk(c, bold, italic, True, False)
        elif tag == "sub":
            for c in node.children:
                walk(c, bold, italic, False, True)
        elif tag == "span":
            classes = node.get("class", [])
            if isinstance(classes, str):
                classes = classes.split()
            if "math-frac" in classes:
                spans = list(node.find_all("span", recursive=False))
                if len(spans) >= 2:
                    num = spans[0].get_text(strip=True)
                    den = spans[1].get_text(strip=True)
                    segments.append(("text", f"({num})/({den})", bold, italic, sup, sub))
                else:
                    for c in node.children:
                        walk(c, bold, italic, sup, sub)
            else:
                for c in node.children:
                    walk(c, bold, italic, sup, sub)
        else:
            for c in node.children:
                walk(c, bold, italic, sup, sub)

    soup = BeautifulSoup(html or "", "html.parser")
    for child in soup.children:
        walk(child)
    while segments and segments[-1][0] == "break":
        segments.pop()

    reordered, group_text, group_imgs = [], [], []
    for seg in segments:
        if seg[0] == "break":
            reordered.extend(group_text)
            reordered.extend(group_imgs)
            reordered.append(seg)
            group_text, group_imgs = [], []
        elif seg[0] == "image_right":
            group_imgs.append(seg)
        else:
            group_text.append(seg)
    reordered.extend(group_text)
    reordered.extend(group_imgs)
    return reordered


def _get_image_bytes(src):
    """Загружает изображение по src (/media/... или http...) и возвращает bytes или None."""
    import os

    import requests as _req
    from django.conf import settings as dj_settings

    try:
        if src.startswith("http://") or src.startswith("https://"):
            r = _req.get(src, timeout=15)
            if r.status_code == 200:
                return r.content
        elif src.startswith("/media/"):
            rel = src[len("/media/") :]
            local = os.path.join(str(getattr(dj_settings, "MEDIA_ROOT", "")), rel.replace("/", os.sep))
            if os.path.exists(local):
                with open(local, "rb") as f:
                    return f.read()
            from django.core.files.storage import default_storage

            try:
                url = default_storage.url(rel)
                if url.startswith("http://") or url.startswith("https://"):
                    r = _req.get(url, timeout=15)
                    if r.status_code == 200:
                        return r.content
                else:
                    with default_storage.open(rel) as f:
                        return f.read()
            except Exception:
                pass
    except Exception as e:
        logger.warning("Не удалось загрузить изображение %s: %s", src, e)
    return None


def _svg_to_png(svg_bytes):
    """Конвертирует SVG байты в PNG байты через PyMuPDF. Возвращает None при ошибке."""
    try:
        import fitz  # PyMuPDF

        fitz_doc = fitz.open("svg", svg_bytes)
        pix = fitz_doc[0].get_pixmap(matrix=fitz.Matrix(3, 3), alpha=False)
        return pix.tobytes("png")
    except Exception as e:
        logger.warning("SVG→PNG конвертация не удалась: %s", e)
        return None


def _image_width(img_data, max_cm=14):
    """Возвращает ширину для вставки изображения: натуральный размер, но не более max_cm."""
    from docx.shared import Cm

    try:
        from PIL import Image as PILImage

        img = PILImage.open(io.BytesIO(img_data))
        w_px = img.size[0]
        info_dpi = img.info.get("dpi", None)
        if info_dpi:
            dpi_x = info_dpi[0] if hasattr(info_dpi, "__getitem__") else info_dpi
        else:
            dpi_x = 96
        if not dpi_x or dpi_x <= 0:
            dpi_x = 96
        w_cm = w_px / dpi_x * 2.54
        return Cm(min(w_cm, max_cm))
    except Exception:
        return Cm(max_cm)


def _render_segments(doc, segments, indent=None, font_size=None, initial_para=None, default_font=None):
    """Рендерит сегменты в документ, создавая параграфы по мере нужды."""
    from docx.shared import Pt

    current = [initial_para]

    def get_para():
        if current[0] is None:
            p = doc.add_paragraph()
            if indent:
                p.paragraph_format.left_indent = indent
            p.paragraph_format.space_after = Pt(1)
            current[0] = p
        return current[0]

    def close():
        current[0] = None

    def _apply(run):
        pass  # font inherited from Normal style

    for seg in segments:
        if seg[0] == "text":
            _, text, bold, italic, sup, sub = seg
            run = get_para().add_run(text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            run.font.superscript = sup
            run.font.subscript = sub
            _apply(run)
        elif seg[0] == "break":
            close()
        elif seg[0] in ("image", "image_right"):
            img_src = seg[1]
            is_right = seg[0] == "image_right"
            img_data = _get_image_bytes(img_src)
            if img_data:
                is_svg = img_src.lower().endswith(".svg") or img_data[:5] in (b"<svg ", b"<?xml")
                if is_svg:
                    img_data = _svg_to_png(img_data)
                if img_data:
                    try:
                        if is_svg:
                            max_svg = 5
                            if is_right:
                                max_svg = min(max_svg, 3)
                            get_para().add_run().add_picture(
                                io.BytesIO(img_data), width=_image_width(img_data, max_cm=max_svg)
                            )
                        elif is_right:
                            get_para().add_run().add_picture(
                                io.BytesIO(img_data), width=_image_width(img_data, max_cm=3)
                            )
                        else:
                            close()
                            doc.add_picture(io.BytesIO(img_data), width=_image_width(img_data, max_cm=4))
                            close()
                    except Exception as e:
                        logger.warning("Не удалось вставить изображение: %s", e)
        elif seg[0] == "table":
            close()
            rows_data = seg[1]
            if not rows_data:
                continue
            max_cols = max(len(row) for row in rows_data)
            if max_cols == 0:
                continue
            tbl = doc.add_table(rows=len(rows_data), cols=max_cols)
            tbl.style = "Table Grid"
            for r_idx, row in enumerate(rows_data):
                for c_idx, cell_html in enumerate(row):
                    if c_idx >= max_cols:
                        break
                    cell = tbl.rows[r_idx].cells[c_idx]
                    cell_segs = _parse_html_segments(cell_html)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    for seg2 in cell_segs:
                        if seg2[0] == "text":
                            run = p.add_run(seg2[1])
                            if seg2[2]:
                                run.bold = True
                            if seg2[3]:
                                run.italic = True
                            run.font.superscript = seg2[4]
                            run.font.subscript = seg2[5]
                            _apply(run)
                        elif seg2[0] == "break":
                            p = cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(1)
            close()


def _extract_task_parts(html):
    """Делит task.text на (ctx_html, body_html).

    Ищет <details class="shared-context"> → возвращает содержимое .shared-context-body
    как ctx_html (или None если нет <details>), а остаток без <details> — как body_html.
    """
    import re

    from bs4 import BeautifulSoup

    if not html or "shared-context" not in html:
        return None, html

    soup = BeautifulSoup(html, "html.parser")
    details = soup.find("details", class_="shared-context")
    if not details:
        return None, html

    body_div = details.find("div", class_="shared-context-body")
    ctx_html = str(body_div) if body_div else ""
    details.decompose()

    body_html = re.sub(r"^(\s*<br\s*/?>)+", "", str(soup)).strip()
    return ctx_html, body_html


def _strip_answer_placeholder(html):
    """Удаляет строки 'Ответ: ___' из HTML задания."""
    import re

    from bs4 import BeautifulSoup

    if not html:
        return html
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup.find_all(["p", "div", "span"]):
        text = tag.get_text()
        if re.search(r"Ответ\s*:\s*[_\s]{3,}", text):
            tag.decompose()
    return str(soup)


def build_variant_docx(variant, include_answers):
    """Строит docx-документ для варианта. Возвращает BytesIO.

    Формат: US Letter, узкие поля — параметры эталонного «33 вариант.docx».
    """
    import requests as _req
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.oxml.ns import qn as _qn
    from docx.shared import Cm, Inches, Pt

    doc = Document()

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)
    rPr = style_normal.element.get_or_add_rPr()
    rFonts = rPr.find(_qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement as _OE

        rFonts = _OE("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(_qn("w:ascii"), "Calibri")
    rFonts.set(_qn("w:hAnsi"), "Calibri")
    rFonts.set(_qn("w:cs"), "Calibri")

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    FS = Pt(11)
    FONT = "Calibri"

    def _sp(p, before=0, after=1):
        if before:
            p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)

    def _run(para, text, bold=False, size=None):
        r = para.add_run(text)
        if size:
            r.font.size = size
        if bold:
            r.bold = True
        return r

    def _shade_cell(cell, fill="D9D9D9"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(title, before=0, after=1)
    _run(title, f"Вариант {variant.number}", bold=True, size=Pt(13))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _sp(sub, before=0, after=3)
    _run(sub, variant.get_exam_type_display(), size=Pt(10))

    tasks = sorted(
        variant.tasks.all(),
        key=lambda t: int(t.number) if t.number.isdigit() else 0,
    )

    printed_ctx_hashes: set = set()
    part1_printed = False
    part2_printed = False

    for task in tasks:
        task_num = int(task.number) if task.number.isdigit() else 0

        if task_num >= 20 and not part2_printed:
            part2_printed = True
            p2 = doc.add_paragraph()
            _sp(p2, before=6, after=2)
            _run(p2, "Часть 2.", bold=True)
        elif not part1_printed and task_num < 20:
            part1_printed = True
            p1 = doc.add_paragraph()
            _sp(p1, before=0, after=2)
            _run(p1, "Часть 1.", bold=True)

        if task.shared_context or task.shared_context_image:
            ctx_key = hash(task.shared_context or "")
            if ctx_key not in printed_ctx_hashes:
                printed_ctx_hashes.add(ctx_key)
                if task.shared_context_image:
                    try:
                        ci_url = task.shared_context_image.url
                        ci_data = (
                            _req.get(ci_url, timeout=15).content
                            if ci_url.startswith("http")
                            else task.shared_context_image.open("rb").read()
                        )
                        doc.add_picture(io.BytesIO(ci_data), width=_image_width(ci_data, max_cm=12))
                    except Exception:
                        logger.warning("Не удалось вставить изображение общего условия")
                if task.shared_context:
                    _render_segments(
                        doc,
                        _parse_html_segments(task.shared_context),
                        font_size=FS,
                        default_font=FONT,
                    )

        ctx_html, body_html = _extract_task_parts(task.text or "")
        if ctx_html:
            ctx_key = hash(ctx_html.strip())
            if ctx_key not in printed_ctx_hashes:
                printed_ctx_hashes.add(ctx_key)
                _render_segments(
                    doc,
                    _parse_html_segments(ctx_html),
                    font_size=FS,
                    default_font=FONT,
                )

        th = doc.add_paragraph()
        _sp(th, before=3, after=1)
        _run(th, f"{task.number}. ", bold=True)

        body_html = _strip_answer_placeholder(body_html)
        if body_html.strip():
            _render_segments(
                doc,
                _parse_html_segments(body_html),
                font_size=FS,
                default_font=FONT,
                initial_para=th,
            )

        if task.image:
            try:
                img_url = task.image.url
                img_data = (
                    _req.get(img_url, timeout=15).content
                    if img_url.startswith("http")
                    else task.image.open("rb").read()
                )
                doc.add_picture(io.BytesIO(img_data), width=_image_width(img_data))
            except Exception:
                logger.warning("Не удалось вставить картинку задания %s", task.number)

    auto_tasks = [t for t in tasks if int(t.number) < 20 if t.number.isdigit()]
    if not auto_tasks:
        auto_tasks = tasks
    if auto_tasks and include_answers:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        sep = doc.add_paragraph()
        _sp(sep, before=0, after=6)
        _run(sep, "Таблица ответов", bold=True, size=Pt(13))

        col_groups = 1
        n = len(auto_tasks)
        rpg = (n + col_groups - 1) // col_groups
        groups = [auto_tasks[g * rpg : (g + 1) * rpg] for g in range(col_groups)]

        tbl = doc.add_table(rows=rpg + 1, cols=col_groups * 2)
        tbl.style = "Table Grid"
        for row in tbl.rows:
            row.cells[0].width = Cm(1)
            row.cells[1].width = Cm(3)

        hdr = tbl.rows[0].cells
        for g in range(col_groups):
            hdr[g * 2].text = "№"
            hdr[g * 2 + 1].text = "Ответ"

        for row_i in range(rpg):
            cells = tbl.rows[row_i + 1].cells
            for g, grp in enumerate(groups):
                if row_i < len(grp):
                    cells[g * 2].text = str(grp[row_i].number)
                    cells[g * 2 + 1].text = grp[row_i].correct_answer or ""

        for row_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _sp(para, before=1, after=1)
                    for run in para.runs:
                        run.font.name = FONT
                        run.font.size = Pt(9)
                        if row_idx == 0 or c_idx % 2 == 0:
                            run.bold = True
                if row_idx == 0 or c_idx % 2 == 0:
                    _shade_cell(cell)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
