import datetime
import logging
import urllib.parse
from functools import wraps

from django.conf import settings as django_settings
from django.contrib.auth import authenticate, login, logout
from django.core.cache import cache
from django.core.paginator import Paginator
from django.db.models import Count, Q
from django.http import HttpResponse
from django.shortcuts import redirect, render
from django.views.decorators.http import require_POST

from .models import Attempt, SchoolClass, Student, Variant

logger = logging.getLogger(__name__)


# ===== ОБЩИЕ ХЕЛПЕРЫ =====


def admin_required(view_func):
    """Декоратор: требует авторизованного админа (Django User с is_staff)."""

    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated or not request.user.is_staff:
            return redirect("admin_login")
        return view_func(request, *args, **kwargs)

    return wrapper


def _safe_int(value, default=0):
    """Безопасное преобразование в int."""
    try:
        return int(value)
    except (ValueError, TypeError):
        return default


def _paginate(request, queryset, per_page=25):
    """Пагинация queryset."""
    paginator = Paginator(queryset, per_page)
    page_number = request.GET.get("page", 1)
    return paginator.get_page(page_number)


def _get_client_ip(request):
    xff = request.META.get("HTTP_X_FORWARDED_FOR", "")
    if xff:
        return xff.split(",")[0].strip()
    return request.META.get("REMOTE_ADDR", "127.0.0.1")


def _check_admin_rate_limit(request):
    ip = _get_client_ip(request)
    lock_key = f"admin_login_lock:{ip}"
    if cache.get(lock_key):
        return "Слишком много попыток. Попробуйте через несколько минут."
    return None


def _record_admin_failed_login(request):
    max_attempts = getattr(django_settings, "LOGIN_MAX_ATTEMPTS", 5)
    cooldown = getattr(django_settings, "LOGIN_COOLDOWN_SECONDS", 300)
    ip = _get_client_ip(request)
    fail_key = f"admin_login_fails:{ip}"
    lock_key = f"admin_login_lock:{ip}"

    fails = cache.get(fail_key, 0) + 1
    cache.set(fail_key, fails, cooldown)
    if fails >= max_attempts:
        cache.set(lock_key, True, cooldown)


# ===== ВХОД / ВЫХОД =====


def admin_login(request):
    if request.user.is_authenticated and request.user.is_staff:
        return redirect("admin_dashboard")

    error = ""
    if request.method == "POST":
        rate_error = _check_admin_rate_limit(request)
        if rate_error:
            error = rate_error
        else:
            username = request.POST.get("username", "").strip()
            password = request.POST.get("password", "").strip()
            user = authenticate(request, username=username, password=password)
            if user and user.is_staff:
                ip = _get_client_ip(request)
                cache.delete(f"admin_login_fails:{ip}")
                cache.delete(f"admin_login_lock:{ip}")
                login(request, user)
                logger.info("Успешный вход админа: %s (IP: %s)", username, ip)
                return redirect("admin_dashboard")
            _record_admin_failed_login(request)
            logger.warning("Неудачная попытка входа админа: %s (IP: %s)", username, _get_client_ip(request))
            error = "Неверный логин или пароль"
    return render(request, "admin/login.html", {"error": error})


@require_POST
def admin_logout(request):
    logout(request)
    return redirect("admin_login")


# ===== ДАШБОРД =====


@admin_required
def dashboard(request):
    stats = {
        "students_count": Student.objects.count(),
        "variants_count": Variant.objects.count(),
        "attempts_count": Attempt.objects.filter(is_finished=True).count(),
        "classes_count": SchoolClass.objects.count(),
    }
    recent_attempts = (
        Attempt.objects.filter(is_finished=True)
        .select_related("student", "variant", "student__school_class")
        .order_by("-finished_at")[:10]
    )

    pending_grading = Attempt.objects.filter(is_finished=True, answers__is_correct=None).distinct().count()

    classes = SchoolClass.objects.filter(is_active=True).annotate(student_count=Count("students"))
    class_stats_list = []
    for sc in classes:
        attempts_qs = Attempt.objects.filter(student__school_class=sc, is_finished=True)
        attempts_count = attempts_qs.count()
        if attempts_count > 0:
            percentages = [
                round(a.score / a.max_score * 100) if a.max_score else 0
                for a in attempts_qs.only("score", "max_score")
            ]
            avg_pct = round(sum(percentages) / len(percentages))
        else:
            avg_pct = None
        class_stats_list.append(
            {"school_class": sc, "attempts_count": attempts_count, "avg_percentage": avg_pct}
        )

    return render(
        request,
        "admin/dashboard.html",
        {
            "stats": stats,
            "recent_attempts": recent_attempts,
            "class_stats": class_stats_list,
            "pending_grading": pending_grading,
        },
    )


# ===== ЭКСПОРТ РЕЗУЛЬТАТОВ =====


def _get_export_filename(class_id, variant_id, ext):
    """Ведомость_9А_вариант_12_18-04-2026.ext"""
    parts = ["Ведомость"]
    if class_id:
        sc = SchoolClass.objects.filter(id=class_id).first()
        if sc:
            parts.append(sc.name)
    if variant_id:
        v = Variant.objects.filter(id=variant_id).first()
        if v:
            parts.append(f"вариант_{v.number}")
    parts.append(datetime.date.today().strftime("%d-%m-%Y"))
    return "_".join(parts) + "." + ext


def _content_disposition(filename):
    """RFC 5987-совместимый заголовок для кириллических имён файлов."""
    safe = urllib.parse.quote(filename, safe="")
    return f"attachment; filename=\"{filename}\"; filename*=UTF-8''{safe}"


def _set_cell_bg(cell, hex_color):
    """Фон ячейки Word через XML (убирает предыдущий shd)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tcPr = cell._tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _task_display(is_correct, awarded_points, max_pts, manual):
    """(текст, hex_цвет|None) для ячейки задания."""
    if manual:
        if is_correct is None:
            return "?", "FFD966"  # ожидает — оранжевый
        pts = awarded_points if awarded_points is not None else 0
        if pts == max_pts:
            return str(pts), "C6EFCE"  # полный балл — зелёный
        if pts > 0:
            return str(pts), "FFEB9C"  # частичный — жёлтый
        return "0", "FFC7CE"  # ноль — красный
    else:
        if is_correct is True:
            return "1", "C6EFCE"
        if is_correct is False:
            return "0", "FFC7CE"
        return "", None


@admin_required
def export_results(request):
    try:
        import openpyxl
        from openpyxl.styles import Alignment, Font, PatternFill
    except ImportError:
        return HttpResponse("openpyxl не установлен. pip install openpyxl", status=500)

    class_id = _safe_int(request.GET.get("class", ""))
    variant_id = _safe_int(request.GET.get("variant", ""))

    attempts = (
        Attempt.objects.filter(is_finished=True)
        .select_related("student", "student__school_class", "variant")
        .annotate(
            pending_count=Count(
                "answers", filter=Q(answers__is_correct=None, answers__task__manual_grading=True)
            )
        )
        .order_by("student__school_class__name", "student__full_name")
    )
    if class_id:
        attempts = attempts.filter(student__school_class_id=class_id)
    if variant_id:
        attempts = attempts.filter(variant_id=variant_id)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Ведомость"

    # --- цвета ---
    hdr_fill = PatternFill("solid", fgColor="1F4E79")
    green_fill = PatternFill("solid", fgColor="C6EFCE")
    yellow_fill = PatternFill("solid", fgColor="FFEB9C")
    red_fill = PatternFill("solid", fgColor="FFC7CE")
    orange_fill = PatternFill("solid", fgColor="FFD966")
    gray_fill = PatternFill("solid", fgColor="D9D9D9")

    # --- шапка документа ---
    info_parts = []
    if class_id:
        sc = SchoolClass.objects.filter(id=class_id).first()
        if sc:
            info_parts.append(f"Класс: {sc.name}   Экзамен: {sc.get_exam_type_display()}")
    if variant_id:
        v = Variant.objects.filter(id=variant_id).first()
        if v:
            info_parts.append(f"Вариант: {v.number}")
    info_parts.append(f"Дата: {datetime.date.today().strftime('%d.%m.%Y')}")

    ws.append(["Ведомость результатов ОГЭ"])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append(["   ".join(info_parts)])
    ws.append([])

    # --- заголовки ---
    headers = [
        "ФИО",
        "Класс",
        "Тип экзамена",
        "Вариант",
        "Дата",
        "Балл",
        "Макс.",
        "Выполнение",
        "Оценка",
        "Статус",
        "Время",
    ]
    ws.append(headers)
    hdr_row_idx = ws.max_row
    for cell in ws[hdr_row_idx]:
        cell.fill = hdr_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(horizontal="center")

    scores, max_scores = [], []

    for a in attempts:
        pct = round(a.score / a.max_score * 100) if a.max_score else 0
        status = "Ожидает" if a.pending_count > 0 else "Проверено"

        ws.append(
            [
                a.student.full_name,
                a.student.school_class.name,
                a.student.school_class.get_exam_type_display(),
                a.variant.number,
                a.finished_at.strftime("%d.%m.%Y %H:%M") if a.finished_at else "",
                a.score,
                a.max_score,
                f"{pct}%",
                a.grade,
                status,
                a.duration_display,
            ]
        )
        r = ws.max_row

        # балл: цвет по оценке
        score_cell = ws.cell(row=r, column=6)
        grade = a.grade
        if grade in ("4", "5"):
            score_cell.fill = green_fill
        elif grade == "3":
            score_cell.fill = yellow_fill
        elif grade == "2":
            score_cell.fill = red_fill

        # статус
        status_cell = ws.cell(row=r, column=10)
        status_cell.fill = orange_fill if status == "Ожидает" else green_fill

        scores.append(a.score)
        max_scores.append(a.max_score)

    # --- итоговая строка ---
    if scores:
        avg_score = round(sum(scores) / len(scores), 1)
        avg_max = max_scores[0] if max_scores else 0
        avg_pct = round(sum(scores) / sum(max_scores) * 100) if sum(max_scores) else 0
        ws.append(["Среднее по классу", "", "", "", "", avg_score, avg_max, f"{avg_pct}%", "", "", ""])
        r = ws.max_row
        for col in range(1, 12):
            c = ws.cell(row=r, column=col)
            c.fill = gray_fill
            c.font = Font(bold=True)

    filename = _get_export_filename(class_id, variant_id, "xlsx")
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    response["Content-Disposition"] = _content_disposition(filename)
    wb.save(response)
    return response


@admin_required
def export_results_docx(request):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor
    except ImportError:
        return HttpResponse("python-docx не установлен.", status=500)

    class_id = _safe_int(request.GET.get("class", ""))
    variant_id = _safe_int(request.GET.get("variant", ""))

    attempts = (
        Attempt.objects.filter(is_finished=True)
        .select_related("student", "student__school_class", "variant")
        .prefetch_related("answers__task")
        .order_by("student__school_class__name", "student__full_name")
    )
    if class_id:
        attempts = attempts.filter(student__school_class_id=class_id)
    if variant_id:
        attempts = attempts.filter(variant_id=variant_id)
    attempts = list(attempts)

    if not attempts:
        return HttpResponse("Нет данных для экспорта.", status=404)

    max_tasks = max((a.variant.tasks.count() for a in attempts), default=25)

    doc = Document()
    section = doc.sections[0]
    section.orientation = 1
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # --- шапка документа ---
    class_name = exam_type_display = variant_num = ""
    if class_id:
        sc = SchoolClass.objects.filter(id=class_id).first()
        if sc:
            class_name = sc.name
            exam_type_display = sc.get_exam_type_display()
    if variant_id:
        v = Variant.objects.filter(id=variant_id).first()
        if v:
            variant_num = v.number
            if not exam_type_display:
                exam_type_display = v.get_exam_type_display()

    title_para = doc.add_heading("Ведомость результатов", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_items = []
    if class_name:
        info_items.append(f"Класс: {class_name}")
    if exam_type_display:
        info_items.append(f"Экзамен: {exam_type_display}")
    if variant_num:
        info_items.append(f"Вариант: {variant_num}")
    info_items.append(f"Дата: {datetime.date.today().strftime('%d.%m.%Y')}")
    info_para = doc.add_paragraph("     ".join(info_items))
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- таблица: ФИО | Дата | Вариант | 1..N | Итого | Оценка | Статус ---
    col_count = 3 + max_tasks + 3
    table = doc.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    hdr[0].text = "ФИО"
    hdr[1].text = "Дата"
    hdr[2].text = "Вариант"
    for i in range(max_tasks):
        hdr[3 + i].text = str(i + 1)
    hdr[3 + max_tasks].text = "Итого"
    hdr[3 + max_tasks + 1].text = "Оценка"
    hdr[3 + max_tasks + 2].text = "Статус"

    for cell in hdr:
        _set_cell_bg(cell, "1F4E79")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    all_scores, all_max_scores = [], []

    for a in attempts:
        # собираем данные ответов
        answers_map = {}
        pending = False
        for ans in a.answers.all():
            if ans.is_correct is None and ans.task.manual_grading:
                pending = True
            answers_map[ans.task.number] = (
                ans.is_correct,
                ans.awarded_points,
                ans.task.points,
                ans.task.manual_grading,
            )

        tasks = list(a.variant.tasks.order_by("id"))
        row_cells = table.add_row().cells

        row_cells[0].text = a.student.full_name
        row_cells[1].text = a.finished_at.strftime("%d.%m.%Y") if a.finished_at else ""
        row_cells[2].text = a.variant.number

        for i, task in enumerate(tasks):
            if i >= max_tasks:
                break
            info = answers_map.get(task.number)
            if info is None:
                text, color = "", None
            else:
                text, color = _task_display(*info)
            cell = row_cells[3 + i]
            cell.text = text
            if color:
                _set_cell_bg(cell, color)

        # Итого
        total_cell = row_cells[3 + max_tasks]
        total_cell.text = str(a.score)
        _set_cell_bg(total_cell, "BDD7EE")

        # Оценка
        row_cells[3 + max_tasks + 1].text = str(a.grade)

        # Статус
        status_cell = row_cells[3 + max_tasks + 2]
        if pending:
            status_cell.text = "Ожидает"
            _set_cell_bg(status_cell, "FFD966")
        else:
            status_cell.text = "✓"
            _set_cell_bg(status_cell, "C6EFCE")

        # шрифт всей строки
        for cell in row_cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size = Pt(8)
        # ФИО — по левому краю
        for para in row_cells[0].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        all_scores.append(a.score)
        all_max_scores.append(a.max_score)

    # --- строка "Среднее" ---
    if all_scores:
        avg_score = round(sum(all_scores) / len(all_scores), 1)
        avg_pct = round(sum(all_scores) / sum(all_max_scores) * 100) if sum(all_max_scores) else 0

        avg_row = table.add_row().cells
        avg_row[0].text = "Среднее по классу"
        avg_row[3 + max_tasks].text = str(avg_score)
        avg_row[3 + max_tasks + 1].text = f"{avg_pct}%"

        for cell in avg_row:
            _set_cell_bg(cell, "D9D9D9")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)
        avg_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # итого и % — выделяем синим
        _set_cell_bg(avg_row[3 + max_tasks], "BDD7EE")
        _set_cell_bg(avg_row[3 + max_tasks + 1], "BDD7EE")

    filename = _get_export_filename(class_id, variant_id, "docx")
    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = _content_disposition(filename)
    doc.save(response)
    return response
